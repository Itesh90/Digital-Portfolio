"""
Resume Parser Service

Handles text extraction from various resume file formats.
"""

import io
from typing import Optional

import pdfplumber
from docx import Document


class ResumeParserService:
    """Service for extracting text from resume files."""
    
    async def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text content from a resume file.
        
        Args:
            file_path: Path to the resume file
            file_type: File extension (.pdf, .docx, .txt)
            
        Returns:
            Extracted text content
        """
        if file_type == ".pdf":
            return await self._extract_pdf(file_path)
        elif file_type in (".docx", ".doc"):
            return await self._extract_docx(file_path)
        elif file_type == ".txt":
            return await self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def extract_text_from_bytes(self, content: bytes, file_type: str) -> str:
        """
        Extract text from file bytes.
        
        Args:
            content: File content as bytes
            file_type: File extension
            
        Returns:
            Extracted text content
        """
        if file_type == ".pdf":
            return self._extract_pdf_bytes(content)
        elif file_type in (".docx", ".doc"):
            return self._extract_docx_bytes(content)
        elif file_type == ".txt":
            return content.decode("utf-8", errors="ignore")
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return "\n\n".join(text_parts)
    
    def _extract_pdf_bytes(self, content: bytes) -> str:
        """Extract text from PDF bytes."""
        text_parts = []
        
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return "\n\n".join(text_parts)
    
    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    
    def _extract_docx_bytes(self, content: bytes) -> str:
        """Extract text from DOCX bytes."""
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    
    async def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
